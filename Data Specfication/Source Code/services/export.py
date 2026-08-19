import os
import json
import zipfile
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EXPORT_DIR = os.path.join(BASE_DIR, "exports")


def ensure_export_dir():
    """Ensures the export storage directory exists."""
    os.makedirs(EXPORT_DIR, exist_ok=True)


def _safe_filename(title, default="Meeting"):
    """Sanitizes meeting title for safe filesystem filenames."""
    clean = "".join(c for c in title if c.isalnum() or c in (' ', '_', '-')).rstrip()
    return clean.replace(' ', '_') or default


def generate_pdf(meeting):
    """Generates a styled, print-ready PDF meeting report using ReportLab."""
    ensure_export_dir()
    clean_title = _safe_filename(meeting.get('title', 'Meeting'))
    filename = f"Meeting_Report_{meeting['id']}_{clean_title}.pdf"
    file_path = os.path.join(EXPORT_DIR, filename)

    doc = SimpleDocTemplate(
        file_path,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    brand_style = ParagraphStyle(
        'BrandTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0f172a")
    )

    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#1e293b"),
        spaceBefore=10,
        spaceAfter=4
    )

    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor("#334155")
    )

    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        textColor=colors.HexColor("#334155"),
        leftIndent=10
    )

    story = [
        Paragraph("<b>TalkToText Pro</b> — Meeting Intelligence Report", brand_style),
        Spacer(1, 3),
        Paragraph(
            f"<b>Meeting:</b> {meeting['title']} | <b>Style:</b> {meeting.get('summary_style', 'Executive Summary')} | <b>Language:</b> {meeting.get('language', 'English')} | <b>Date:</b> {meeting.get('created_at', 'N/A')}",
            body_style
        ),
        Spacer(1, 6),
        HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#cbd5e1"), spaceBefore=2, spaceAfter=8),
        Paragraph(f"{meeting.get('summary_style', 'Executive Summary')}", h2_style),
        Paragraph(meeting.get('summary', 'No summary available.'), body_style),
        Spacer(1, 8)
    ]

    key_points = meeting.get('key_points', [])
    if key_points:
        story.append(Paragraph("Key Discussion Points", h2_style))
        for pt in key_points:
            story.append(Paragraph(f"• {pt}", bullet_style))
            story.append(Spacer(1, 2))
        story.append(Spacer(1, 6))

    decisions = meeting.get('decisions', [])
    if decisions:
        story.append(Paragraph("Decisions Made", h2_style))
        for dec in decisions:
            story.append(Paragraph(f"✓ {dec}", bullet_style))
            story.append(Spacer(1, 2))
        story.append(Spacer(1, 6))

    action_items = meeting.get('action_items', [])
    if action_items:
        story.append(Paragraph("Action Items & Deliverables", h2_style))
        table_data = [["Assignee", "Task Description", "Deadline", "Priority", "Status"]]
        for item in action_items:
            table_data.append([
                item.get('assignee', 'Unassigned'),
                item.get('task', ''),
                item.get('deadline', 'TBD'),
                item.get('priority', 'Medium'),
                item.get('status', 'Pending')
            ])

        table = Table(table_data, colWidths=[1.1 * inch, 3.2 * inch, 1.0 * inch, 0.9 * inch, 0.9 * inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 5),
            ('TOPPADDING', (0, 0), (-1, 0), 5),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8.5),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")])
        ]))
        story.append(table)
        story.append(Spacer(1, 8))

    sentiment = meeting.get('sentiment', {})
    if sentiment:
        story.append(Paragraph("Sentiment & Collaboration Assessment", h2_style))
        s_text = (
            f"<b>Overall Tone:</b> {sentiment.get('tone', 'Constructive')} | "
            f"<b>Positive:</b> {sentiment.get('positive_pct', 75)}% | "
            f"<b>Neutral:</b> {sentiment.get('neutral_pct', 20)}% | "
            f"<b>Negative:</b> {sentiment.get('negative_pct', 5)}%<br/>"
            f"<b>Insights:</b> {sentiment.get('insights', '')}"
        )
        story.append(Paragraph(s_text, body_style))
        story.append(Spacer(1, 8))

    speaker_transcript = meeting.get('speaker_transcript', [])
    if speaker_transcript:
        story.append(Paragraph("Speaker Identification & Dialogue", h2_style))
        for spk in speaker_transcript:
            story.append(Paragraph(f"<b>{spk.get('speaker', 'Speaker')}:</b> {spk.get('text', '')}", body_style))
            story.append(Spacer(1, 3))

    doc.build(story)
    return file_path


def generate_docx(meeting):
    """Generates a formatted Microsoft Word (.docx) document."""
    ensure_export_dir()
    clean_title = _safe_filename(meeting.get('title', 'Meeting'))
    filename = f"Meeting_Report_{meeting['id']}_{clean_title}.docx"
    file_path = os.path.join(EXPORT_DIR, filename)

    doc = Document()
    doc.add_heading('TalkToText Pro — Meeting Intelligence Report', level=0)

    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Meeting Title: {meeting['title']}\n").bold = True
    meta_p.add_run(f"Style: {meeting.get('summary_style', 'Executive Summary')} | Language: {meeting.get('language', 'English')} | Date: {meeting.get('created_at', 'N/A')}")

    doc.add_heading(f"{meeting.get('summary_style', 'Executive Summary')}", level=1)
    doc.add_paragraph(meeting.get('summary', 'No summary available.'))

    key_points = meeting.get('key_points', [])
    if key_points:
        doc.add_heading('Key Discussion Points', level=1)
        for pt in key_points:
            doc.add_paragraph(pt, style='List Bullet')

    decisions = meeting.get('decisions', [])
    if decisions:
        doc.add_heading('Decisions Made', level=1)
        for dec in decisions:
            doc.add_paragraph(dec, style='List Bullet')

    action_items = meeting.get('action_items', [])
    if action_items:
        doc.add_heading('Action Items & Deliverables', level=1)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Assignee'
        hdr_cells[1].text = 'Task Description'
        hdr_cells[2].text = 'Deadline'
        hdr_cells[3].text = 'Priority'
        hdr_cells[4].text = 'Status'

        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('assignee', 'Unassigned')
            row_cells[1].text = item.get('task', '')
            row_cells[2].text = item.get('deadline', 'TBD')
            row_cells[3].text = item.get('priority', 'Medium')
            row_cells[4].text = item.get('status', 'Pending')

    sentiment = meeting.get('sentiment', {})
    if sentiment:
        doc.add_heading('Sentiment Analysis', level=1)
        sp = doc.add_paragraph()
        sp.add_run(f"Tone: {sentiment.get('tone', 'Constructive')}\n")
        sp.add_run(f"Distribution: Positive {sentiment.get('positive_pct', 75)}%, Neutral {sentiment.get('neutral_pct', 20)}%, Negative {sentiment.get('negative_pct', 5)}%\n")
        sp.add_run(f"Observations: {sentiment.get('insights', '')}")

    speaker_transcript = meeting.get('speaker_transcript', [])
    if speaker_transcript:
        doc.add_heading('Speaker Breakdown', level=1)
        for spk in speaker_transcript:
            p = doc.add_paragraph()
            p.add_run(f"{spk.get('speaker', 'Speaker')}: ").bold = True
            p.add_run(spk.get('text', ''))

    doc.save(file_path)
    return file_path


def generate_txt(meeting):
    """Generates clean plaintext meeting notes."""
    ensure_export_dir()
    clean_title = _safe_filename(meeting.get('title', 'Meeting'))
    filename = f"Meeting_Notes_{meeting['id']}_{clean_title}.txt"
    file_path = os.path.join(EXPORT_DIR, filename)

    lines = [
        "=" * 60,
        "TALKTOTEXT PRO — MEETING INTELLIGENCE REPORT",
        "=" * 60,
        f"Title:    {meeting.get('title', 'Untitled')}",
        f"Date:     {meeting.get('created_at', 'N/A')}",
        f"Language: {meeting.get('language', 'Auto-Detect')}",
        f"Duration: {int(meeting.get('duration_seconds', 0) // 60)} minutes",
        "-" * 60,
        "\n--- EXECUTIVE SUMMARY ---\n",
        meeting.get('summary', 'No summary available.'),
        "\n--- KEY DISCUSSION POINTS ---\n"
    ]

    for pt in meeting.get('key_points', []):
        lines.append(f"• {pt}")

    lines.append("\n--- AGREED DECISIONS ---\n")
    for dec in meeting.get('decisions', []):
        lines.append(f"✓ {dec}")

    lines.append("\n--- ACTION ITEMS & DELIVERABLES ---\n")
    for item in meeting.get('action_items', []):
        lines.append(f"[{item.get('status', 'Pending')}] {item.get('task', '')} (Assignee: {item.get('assignee', 'Unassigned')}, Due: {item.get('deadline', 'TBD')})")

    lines.append("\n--- SPEAKER TRANSCRIPT ---\n")
    for spk in meeting.get('speaker_transcript', []):
        lines.append(f"{spk.get('speaker', 'Speaker')}: {spk.get('text', '')}\n")

    with open(file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return file_path


def generate_ics(meeting):
    """Generates standard iCalendar (.ics) file with events for deliverables."""
    ensure_export_dir()
    clean_title = _safe_filename(meeting.get('title', 'Meeting'))
    filename = f"Action_Items_{meeting['id']}_{clean_title}.ics"
    file_path = os.path.join(EXPORT_DIR, filename)

    now_str = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    ics_lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//TalkToText Pro//Meeting Deliverables//EN",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH"
    ]

    action_items = meeting.get('action_items', [])
    if not action_items:
        ics_lines.extend([
            "BEGIN:VEVENT",
            f"UID:meeting-{meeting['id']}-recap@talktotext.pro",
            f"DTSTAMP:{now_str}",
            f"DTSTART:{now_str}",
            f"DTEND:{now_str}",
            f"SUMMARY:Recap: {meeting.get('title', 'Meeting')}",
            f"DESCRIPTION:{meeting.get('summary', '')[:200]}",
            "STATUS:CONFIRMED",
            "END:VEVENT"
        ])
    else:
        for idx, item in enumerate(action_items, start=1):
            task_desc = item.get('task', 'Deliverable')
            assignee = item.get('assignee', 'Team')
            deadline = item.get('deadline', 'Upcoming')
            ics_lines.extend([
                "BEGIN:VEVENT",
                f"UID:task-{meeting['id']}-{idx}@talktotext.pro",
                f"DTSTAMP:{now_str}",
                f"DTSTART:{now_str}",
                f"DTEND:{now_str}",
                f"SUMMARY:[Action Item] {task_desc} ({assignee})",
                f"DESCRIPTION:Task: {task_desc}\\nAssignee: {assignee}\\nDue: {deadline}\\nMeeting: {meeting.get('title', '')}",
                "STATUS:CONFIRMED",
                "BEGIN:VALARM",
                "TRIGGER:-PT15M",
                "ACTION:DISPLAY",
                f"DESCRIPTION:Reminder: {task_desc}",
                "END:VALARM",
                "END:VEVENT"
            ])

    ics_lines.append("END:VCALENDAR")

    with open(file_path, "w", encoding="utf-8") as f:
        f.write("\r\n".join(ics_lines))

    return file_path


def generate_bulk_zip(meetings_list, formats=None):
    """Exports multiple meetings into a consolidated ZIP archive."""
    ensure_export_dir()
    if not formats:
        formats = ['pdf', 'txt']

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_filename = f"TalkToText_Bulk_Export_{timestamp}.zip"
    zip_path = os.path.join(EXPORT_DIR, zip_filename)

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        readme_content = f"# TalkToText Pro — Bulk Meeting Export\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\nTotal Meetings: {len(meetings_list)}\n\n"

        for m in meetings_list:
            clean_title = _safe_filename(m.get('title', 'Meeting'))
            m_id = m.get('id', 'item')
            readme_content += f"- Meeting ID {m_id}: {m.get('title', 'Untitled')} ({m.get('created_at', '')[:10]})\n"

            if 'pdf' in formats:
                try:
                    pdf_path = generate_pdf(m)
                    if os.path.exists(pdf_path):
                        zipf.write(pdf_path, arcname=f"PDFs/{os.path.basename(pdf_path)}")
                except Exception as err:
                    print(f"[Bulk export PDF notice for meeting {m_id}]: {err}")

            if 'docx' in formats:
                try:
                    docx_path = generate_docx(m)
                    if os.path.exists(docx_path):
                        zipf.write(docx_path, arcname=f"DOCX/{os.path.basename(docx_path)}")
                except Exception as err:
                    print(f"[Bulk export DOCX notice for meeting {m_id}]: {err}")

            if 'txt' in formats:
                try:
                    txt_path = generate_txt(m)
                    if os.path.exists(txt_path):
                        zipf.write(txt_path, arcname=f"TXT/{os.path.basename(txt_path)}")
                except Exception as err:
                    print(f"[Bulk export TXT notice for meeting {m_id}]: {err}")

            if 'md' in formats or 'markdown' in formats:
                md_content = f"# {m.get('title', 'Meeting Notes')}\n\n"
                md_content += f"**Date:** {m.get('created_at', '')} | **Language:** {m.get('language', 'EN')}\n\n"
                md_content += f"## Executive Summary\n{m.get('summary', 'No summary available.')}\n\n"
                md_content += "## Key Discussion Points\n"
                for kp in m.get('key_points', []):
                    md_content += f"- {kp}\n"
                md_content += "\n## Decisions Made\n"
                for dec in m.get('decisions', []):
                    md_content += f"- {dec}\n"
                md_content += "\n## Action Items\n"
                for ai in m.get('action_items', []):
                    md_content += f"- [ ] **{ai.get('task', '')}** (Assignee: {ai.get('assignee', 'TBD')}, Due: {ai.get('deadline', 'TBD')})\n"
                md_filename = f"Meeting_{m_id}_{clean_title}.md"
                zipf.writestr(f"Markdown/{md_filename}", md_content)

            if 'json' in formats:
                json_str = json.dumps(m, indent=2, default=str)
                zipf.writestr(f"JSON/Meeting_{m_id}_{clean_title}.json", json_str)

        zipf.writestr("INDEX.txt", readme_content)

    return zip_path
